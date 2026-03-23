"""
Import 84 cases with images, tables, and structured sections.
Extracts images to public/cases/, parses metadata from tables.
"""
import os
import json
import re
from docx import Document

CASES_DIR = os.path.expanduser("~/Desktop/AI/casebooks/cases/")
OUTPUT_DIR = os.path.expanduser("~/case-partner/public/cases/")
os.makedirs(OUTPUT_DIR, exist_ok=True)

def slug(filename):
    return filename.replace(".docx", "").replace(" ", "_")

def extract_images(doc, case_slug):
    """Extract images and return list of paths."""
    images = []
    img_rels = [r for r in doc.part.rels.values() if "image" in r.reltype]
    for i, rel in enumerate(img_rels):
        ext = rel.target_ref.split(".")[-1]
        fname = f"{case_slug}_{i+1}.{ext}"
        path = os.path.join(OUTPUT_DIR, fname)
        with open(path, "wb") as f:
            f.write(rel.target_part.blob)
        images.append(f"/cases/{fname}")
    return images

def extract_table_metadata(doc):
    """Extract metadata from the first table."""
    meta = {}
    if doc.tables:
        for row in doc.tables[0].rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                key = cells[0]
                val = cells[1]
                if "类型" in key: meta["type"] = val
                elif "难度" in key: meta["difficulty"] = val
                elif "行业" in key: meta["industry"] = val
                elif "来源" in key: meta["source"] = val
                elif "形式" in key: meta["format"] = val
    return meta

def parse_case(filepath):
    doc = Document(filepath)
    case_slug = slug(os.path.basename(filepath))

    # Extract images
    images = extract_images(doc, case_slug)

    # Extract table metadata
    meta = extract_table_metadata(doc)

    # Parse paragraphs by style/section
    texts = []
    for p in doc.paragraphs:
        if p.text.strip():
            texts.append({
                "text": p.text.strip(),
                "style": p.style.name,
            })

    title = texts[0]["text"] if texts else case_slug

    # Split title
    title_parts = re.split(r"\s*[—–]\s*", title, maxsplit=1)
    title_en = title_parts[0].strip()
    title_cn = title_parts[1].strip() if len(title_parts) > 1 else ""

    # Section parsing
    prompt_en = ""
    prompt_cn = ""
    clarifying_lines = []
    framework_lines = []
    exhibit1_lines = []
    exhibit2_lines = []
    qa_lines = []
    recommendation_lines = []
    tips_lines = []
    deep_lines = []

    section = "intro"
    for item in texts:
        t = item["text"]
        style = item["style"]

        # Skip separators and metadata
        if t.startswith("──") or t == "---" or t == "简要信息" or t == "Prompt":
            continue

        # Section detection
        if "English" in t and "原文" in t:
            section = "prompt_en"; continue
        elif "中文翻译" in t:
            section = "prompt_cn"; continue
        elif "澄清信息" in t or "Clarifying" in t:
            section = "clarifying"; continue
        elif t == "Framework":
            section = "framework"; continue
        elif "以下是框架的详细文字拆解" in t:
            section = "framework_detail"; continue
        elif "本Case的分析框架" in t:
            section = "framework"; continue
        elif "Exhibit #1" in t and "Heading" in style:
            section = "exhibit1"; continue
        elif "Exhibit #2" in t and "Heading" in style:
            section = "exhibit2"; continue
        elif "问题 & 回答" in t or "问题 &" in t:
            section = "qa"; continue
        elif t.startswith("Q1:") or t.startswith("Q2:") or t.startswith("Q3:"):
            section = "qa"
        elif t == "Recommendation" or ("Recommendation" in t and "Heading" in style):
            section = "recommendation"; continue
        elif "面试技巧" in t or "得分亮点" in t or "Case难点" in t:
            section = "tips"
        elif "深度分析" in t or "深度问答" in t:
            section = "deep"
        elif "分析思路补充" in t:
            section = "framework_detail"
        elif t == "总结" and "Heading" in style:
            section = "tips"; continue

        # Collect
        if section == "prompt_en" and not prompt_en:
            prompt_en = t
        elif section == "prompt_cn" and not prompt_cn:
            prompt_cn = re.sub(r"\s*─+\s*$", "", t)
        elif section == "clarifying":
            if not t.startswith("面试提示") and not t.startswith("如果面试"):
                clarifying_lines.append(t)
        elif section in ("framework", "framework_detail"):
            # Use heading style to add structure
            if "Heading 3" in style:
                framework_lines.append(f"\n**{t}**")
            else:
                framework_lines.append(t)
        elif section == "exhibit1":
            if "Heading" in style:
                exhibit1_lines.append(f"\n**{t}**")
            else:
                exhibit1_lines.append(t)
        elif section == "exhibit2":
            if "Heading" in style:
                exhibit2_lines.append(f"\n**{t}**")
            else:
                exhibit2_lines.append(t)
        elif section == "qa":
            if t.startswith("Q"):
                qa_lines.append(f"\n**{t}**")
            else:
                qa_lines.append(t)
        elif section == "recommendation":
            recommendation_lines.append(t)
        elif section == "tips":
            if "Heading" in style or "难点" in t or "亮点" in t or "技巧" in t:
                tips_lines.append(f"\n**{t}**")
            else:
                tips_lines.append(f"• {t}" if "Bullet" in style else t)
        elif section == "deep":
            if t.startswith("Q:"):
                deep_lines.append(f"\n**{t}**")
            elif "深度" in t and "Heading" in style:
                deep_lines.append(f"\n**{t}**")
            else:
                deep_lines.append(t)

    # Clean up framework - deduplicate heading+content pairs
    clean_framework = []
    prev_heading = ""
    for line in framework_lines:
        if line.startswith("\n**"):
            heading = line.strip("*\n ")
            if heading != prev_heading:
                clean_framework.append(line)
                prev_heading = heading
        else:
            clean_framework.append(f"  {line}")
            prev_heading = ""

    # Map difficulty
    diff_map = {"简单": "Easy", "中等": "Medium", "困难": "Hard", "难": "Hard"}
    difficulty = meta.get("difficulty", "Medium")
    difficulty = diff_map.get(difficulty, difficulty)

    # Map type from table
    type_map = {
        "Profitability": "Profitability",
        "Market Entry": "Market Entry",
        "Market Sizing": "Market Sizing",
        "Pricing": "Pricing",
        "M&A": "M&A",
        "Growth Strategy": "Growth Strategy",
        "Operations": "Operations",
    }
    case_type = meta.get("type", "Profitability")
    # Try to match
    for k, v in type_map.items():
        if k.lower() in case_type.lower():
            case_type = v
            break

    # Industry from table
    industry = meta.get("industry", "General")
    # Clean up industry
    if "（" in industry:
        industry = industry.split("（")[0].strip()

    sections = {}
    if clarifying_lines:
        sections["clarifying"] = "\n".join(clarifying_lines)
    if clean_framework:
        sections["framework"] = "\n".join(clean_framework)
    if exhibit1_lines:
        sections["exhibit1"] = "\n".join(exhibit1_lines)
    if exhibit2_lines:
        sections["exhibit2"] = "\n".join(exhibit2_lines)
    if qa_lines:
        sections["qa"] = "\n".join(qa_lines)
    if recommendation_lines:
        sections["recommendation"] = "\n".join(recommendation_lines)
    if tips_lines:
        sections["tips"] = "\n".join(tips_lines)
    if deep_lines:
        sections["deep_analysis"] = "\n".join(deep_lines)

    # Image assignments (convention: image1=issue tree, image2=exhibit1, image3=exhibit2)
    if len(images) >= 1:
        sections["issue_tree_img"] = images[0]
    if len(images) >= 2:
        sections["exhibit1_img"] = images[1]
    if len(images) >= 3:
        sections["exhibit2_img"] = images[2]

    return {
        "title": title_en,
        "title_cn": title_cn,
        "type": case_type,
        "difficulty": difficulty,
        "industry": industry,
        "source": meta.get("source", "LemonTalk CaseBook 2026"),
        "prompt_en": prompt_en,
        "prompt_cn": prompt_cn,
        "sections": sections,
        "tags": [case_type, industry],
    }


# Parse all
cases = []
for f in sorted(os.listdir(CASES_DIR)):
    if not f.endswith(".docx"): continue
    try:
        c = parse_case(os.path.join(CASES_DIR, f))
        cases.append(c)
    except Exception as e:
        print(f"ERROR {f}: {e}")

print(f"Parsed {len(cases)} cases")
print(f"Images extracted to {OUTPUT_DIR}")

# Count images
img_count = len([f for f in os.listdir(OUTPUT_DIR) if f.endswith('.png')])
print(f"Total images: {img_count}")

# Output TypeScript
ts = """// Auto-generated from LemonTalk CaseBook 2026
// 84 structured cases with images and sections

export interface CaseSection {
  clarifying?: string;
  framework?: string;
  exhibit1?: string;
  exhibit2?: string;
  qa?: string;
  recommendation?: string;
  tips?: string;
  deep_analysis?: string;
  issue_tree_img?: string;
  exhibit1_img?: string;
  exhibit2_img?: string;
}

export interface CaseData {
  title: string;
  title_cn: string;
  type: string;
  difficulty: string;
  industry: string;
  source: string;
  prompt_en: string;
  prompt_cn: string;
  sections: CaseSection;
  tags: string[];
}

export const CASES_DATA: CaseData[] = """

ts += json.dumps(cases, ensure_ascii=False, indent=2)
ts += ";\n"

output_path = os.path.expanduser("~/case-partner/lib/cases_data.ts")
with open(output_path, "w") as f:
    f.write(ts)

print(f"\nWritten to {output_path}")

# Stats
types = {}
industries = {}
diffs = {}
for c in cases:
    types[c["type"]] = types.get(c["type"], 0) + 1
    industries[c["industry"]] = industries.get(c["industry"], 0) + 1
    diffs[c["difficulty"]] = diffs.get(c["difficulty"], 0) + 1

print("\nBy type:")
for t, n in sorted(types.items(), key=lambda x: -x[1]):
    print(f"  {t}: {n}")
print("\nBy difficulty:")
for d, n in sorted(diffs.items()):
    print(f"  {d}: {n}")
print("\nBy industry:")
for i, n in sorted(industries.items(), key=lambda x: -x[1])[:10]:
    print(f"  {i}: {n}")
